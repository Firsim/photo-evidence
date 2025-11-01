# ФОТОТАБЛИЦА v1.0
# Генератор документа Word из фото (JPEG/HEIC) с датой, местом и моделью камеры.
# Создаёт отчёт для суда: !_ФОТОТАБЛИЦА_дата_время.docx

import os
import sys
import datetime
from PIL import Image
from PIL.ExifTags import TAGS, GPSTAGS
from docx import Document
from docx.shared import Inches
from geopy.geocoders import Nominatim
from io import BytesIO
import pillow_heif  # Для поддержки HEIC
import time


# === НАСТРОЙКИ ===
USER_AGENT = "court_photo_evidence"
OUTPUT_PREFIX = "!_ФОТОТАБЛИЦА_"  # Префикс выходного файла
SUPPORTED_FORMATS = (".jpg", ".jpeg", ".JPG", ".JPEG", ".heic", ".HEIF", ".heif")


# Инициализация геокодера
geolocator = Nominatim(user_agent=USER_AGENT)
# Регистрация HEIC как формата для Pillow
pillow_heif.register_heif_opener()


def get_gps_location(gps_data):
    """Преобразует GPS-координаты из EXIF в десятичные градусы."""
    if not gps_data or "GPSLatitude" not in gps_data or "GPSLongitude" not in gps_data:
        return None

    def convert_to_degrees(value):
        d, m, s = value
        return d + (m / 60.0) + (s / 3600.0)

    try:
        lat = convert_to_degrees(gps_data["GPSLatitude"])
        lon = convert_to_degrees(gps_data["GPSLongitude"])

        if gps_data.get("GPSLatitudeRef") != "N":
            lat = -lat
        if gps_data.get("GPSLongitudeRef") != "E":
            lon = -lon

        return round(lat, 6), round(lon, 6)
    except Exception as e:
        print(f"  ⚠️ Ошибка обработки GPS: {e}")
        return None


def get_address_from_coords(lat, lon):
    """Получение читаемого адреса по координатам."""
    try:
        location = geolocator.reverse((lat, lon), language="ru", timeout=10)
        return location.address
    except Exception as e:
        return f"{lat:.6f}, {lon:.6f}"


def extract_image_info(image_path):
    """
    Извлекает из фото:
    - дату съёмки
    - координаты и адрес
    - модель камеры
    Возвращает словарь или None.
    """
    filename = os.path.basename(image_path)
    print(f"  → Обработка: {filename}...", end="", flush=True)

    try:
        image = Image.open(image_path)
    except Exception as e:
        print(f" ОШИБКА: не удалось открыть файл.")
        return None

    exif = image.getexif()
    if not exif:
        print(" НЕТ EXIF.")
        return None

    # Извлечение тегов
    exif_data = {}
    for tag_id in exif:
        tag = TAGS.get(tag_id, tag_id)
        value = exif.get(tag_id)
        if isinstance(value, bytes):
            value = value.decode('utf-8', errors='ignore')
        exif_data[tag] = value

    # GPS данные
    try:
        gps_exif = exif.get_ifd(0x8825)
        gps_data = {GPSTAGS.get(key, key): val for key, val in gps_exif.items()}
    except Exception:
        gps_data = {}

    # Дата съёмки
    date_str = exif_data.get("DateTimeOriginal") or exif_data.get("DateTime")
    if not date_str:
        print(" НЕТ ДАТЫ.")
        return None

    try:
        date_taken = datetime.datetime.strptime(date_str, "%Y:%m:%d %H:%M:%S")
    except Exception:
        print(" ОШИБКА ДАТЫ.")
        return None

    # Координаты
    coords = get_gps_location(gps_data)
    address = get_address_from_coords(*coords) if coords else "Координаты недоступны"

    # Модель камеры
    camera_model = exif_data.get("Model", "Модель неизвестна")

    print(" ГОТОВО.")

    return {
        "image": image.copy(),
        "date_taken": date_taken,
        "address": address,
        "camera_model": camera_model,
    }


def create_document(photos_info, output_dir):
    """Создаёт документ Word в указанной папке."""
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d_%H-%M")
    output_path = os.path.join(output_dir, f"{OUTPUT_PREFIX}{timestamp}.docx")

    print(f"\n📝 Создаём документ: {os.path.basename(output_path)}...")

    doc = Document()
    doc.add_heading('ФОТОТАБЛИЦА', 0)

    # Сортировка: от старых к новым
    photos_info.sort(key=lambda x: x["date_taken"])

    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Фотография'
    hdr_cells[1].text = 'Дата съёмки\n(и модель камеры)'
    hdr_cells[2].text = 'Место'

    for i, info in enumerate(photos_info, 1):
        row_cells = table.add_row().cells
        print(f"  ➕ Добавлено фото {i}/{len(photos_info)}")

        # Миниатюра
        run = row_cells[0].paragraphs[0].add_run()
        try:
            img_to_save = info["image"].convert("RGB") if info["image"].mode != "RGB" else info["image"]
            img_byte_arr = BytesIO()
            img_to_save.save(img_byte_arr, format='JPEG', quality=85)
            img_byte_arr.seek(0)
            run.add_picture(img_byte_arr, width=Inches(1.5))
        except Exception:
            row_cells[0].text = "Ошибка загрузки"

        # Дата + камера
        date_str = info["date_taken"].strftime("%d.%m.%Y %H:%M:%S")
        row_cells[1].text = f"Дата: {date_str}\nКамера: {info['camera_model']}"

        # Адрес
        row_cells[2].text = info["address"]

        time.sleep(0.5)  # Лёгкая задержка для плавности

    doc.save(output_path)
    print(f"\n✅ УСПЕХ: документ сохранён\n    {output_path}")
    return output_path


def main(folder_path):
    """Основная логика."""
    if not os.path.exists(folder_path):
        print(f"❌ ОШИБКА: папка не найдена — {folder_path}")
        return

    if not os.path.isdir(folder_path):
        print(f"❌ ОШИБКА: путь не является папкой — {folder_path}")
        return

    print(f"\n🔍 Поиск фото в: {folder_path}")
    photos = [
        os.path.join(folder_path, f) for f in os.listdir(folder_path)
        if f.lower().endswith(SUPPORTED_FORMATS)
    ]

    if not photos:
        print(f"❌ Нет подходящих фото в папке.\n   Поддерживаются: {SUPPORTED_FORMATS}")
        return

    print(f"📦 Найдено {len(photos)} фото. Начинаем обработку...\n")
    photo_info_list = []

    for photo in photos:
        info = extract_image_info(photo)
        if info:
            photo_info_list.append(info)

    if not photo_info_list:
        print("\n❌ Ни одно фото не содержит необходимых данных (EXIF, дата, GPS).")
        return

    print(f"\n📊 Обработано: {len(photo_info_list)} фото из {len(photos)}")
    create_document(photo_info_list, folder_path)


# === ЗАПУСК ИЗ КОМАНДНОЙ СТРОКИ ===
if __name__ == "__main__":
    print("🚀 ФОТОТАБЛИЦА v1.0 — подготовка доказательств для суда")

    # Если путь передан как аргумент (из .bat)
    if len(sys.argv) > 1:
        folder = sys.argv[1].strip('"')  # Убираем кавычки, если были
    else:
        # Или запросим вручную
        folder = input("\n📁 Введите путь к папке с фотографиями: ").strip().strip('"')

    main(folder)

    print("\n🔚 Работа завершена. Нажмите Enter, чтобы выйти...")
    input()  # Задержка перед закрытием
